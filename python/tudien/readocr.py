# -*- coding: utf-8 -*-
"""
Created on Tue Apr 24 13:15:35 2018

@author: dang

Written for Python 2.7
"""

from __future__ import unicode_literals
import docx

doc = docx.Document('TudienKHCN1.ocr.docx')
#for p in doc.paragraphs:
#    print(len(p.runs))

def read_format(word_style_names, default_bold, default_italic):
    word_format_bolds = [default_bold for item in word_style_names]
    word_format_italics = [default_italic for item in word_style_names]
    for k in range(len(word_style_names)):
        if 'Not Bold' in word_style_names[k]:
            word_format_bolds[k] = False
        elif 'Bold' in word_style_names[k]:
            word_format_bolds[k] = True
        if 'Not Italic' in word_style_names[k]:
            word_format_italics[k] = False
        elif 'Bold' in word_style_names[k]:
            word_format_italics[k] = True
    return word_format_bolds, word_format_italics

def check_capital(phrase):
    format_capital = []
    words = phrase.split()
    for k in range(len(words)):
        if words[k].isupper():                
            format_capital.append(True)
        else:
            format_capital.append(False)
    return words, format_capital

def merge_similar_series(words, case_capital):
    newwords = list(words)
    newcase_capital = list(case_capital)
    for k in range(len(words)-1,0,-1):
        if newcase_capital[k] == newcase_capital[k-1]:
            newwords[k-1] = ' '.join([newwords[k-1], newwords[k]])
            del newwords[k]
            del newcase_capital[k]
    return newwords, newcase_capital
   
def re_parse(word_texts, word_format_bolds, word_format_italics):
    """Fix the leftover mistake in the data parsed, some example items:
    TOÁN, S_CHÊ figure 
    hình 
    (bản vẽ) - is a separate item, in fact it is a comment for previous item
    AB-Betrieb m Đ_TỬ dass AB mode chê'độ hạng AB, splitted into:
    ['AB-Betrieb', 'm', 'Đ_TỬ', 'dass AB', 'mode', 'chê'độ hạng', 'AB']
    """
    newword_texts = []
    newwordcase_capitals = []
    newword_format_bolds = []
    newword_format_italics = []
    newlinerole_comment = []
    for k in range(len(word_texts)):
        phrase = word_texts[k]
        words, wordcase_capital = check_capital(phrase)
        words, wordcase_capital = merge_similar_series(words, wordcase_capital)
        newword_texts.extend(words)
        newwordcase_capitals.extend(wordcase_capital)
        for n in range(len(words)):
            newword_format_bolds.append(word_format_bolds[k])
            newword_format_italics.append(word_format_italics[k])
            # if the word is in (), it should be a comment for the previous one
            if words[n].strip(',')[0]=='(' and words[n].strip(',')[-1]==')':
                newlinerole_comment.append(True)
            else:
                newlinerole_comment.append(False)
    return newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics, newlinerole_comment


def analyze_line(word_texts, wordcase_capitals, word_format_bolds, word_format_italics, linerole_comment):
    """Algorithm to analyze a line:
    A line needs at least 4 words. If not,  neglect the line.
    First word is considered the main word, in German.
    Second word is considered the type of the German word.
    Third word must be CAPITAL, indicating the FIELD. If not, neglect the line.
    After a CAPITAL word: lowercase non-italic words are English, lowercase italic words are Vietnamese.
    Split items for different FIELDS.
    Split items for new English meaning (the precedent word is Vietnamese).
    """
    de_word = []
    en_word = []
    vi_word = []
    type_word = []
    field_word = []
    if len(word_texts) < 4:
        out_message = 'Line has less than 4 elements, skipped.'
        return de_word, type_word, field_word, en_word, vi_word, out_message
    if wordcase_capitals[2] != True:
        out_message = '3rd word is not CAPITAL, syntax not compatible, skipped.'
        return de_word, type_word, field_word, en_word, vi_word, out_message
    item = 0
    de_word.append(word_texts[0])
    type_word.append(word_texts[1])
    field_word.append(word_texts[2])
    en_word.append('')
    vi_word.append('')
    for k in range(3, len(word_texts)):
        if wordcase_capitals[k] == True:
            item += 1
            de_word.append(word_texts[0])
            type_word.append(word_texts[1])
            field_word.append(word_texts[k])
            en_word.append('')
            vi_word.append('')
        else:
            if word_format_italics[k] != True:
                if word_format_italics[k-1] == True and wordcase_capitals[k-1] != True:
                    item += 1
                    de_word.append(word_texts[0])
                    type_word.append(word_texts[1])
                    field_word.append(field_word[item-1])
                    en_word.append('')
                    en_word[item] += ' '+word_texts[k]
                    vi_word.append('')
                else:
                    en_word[item] += ' '+word_texts[k]
            else:
                vi_word[item] += ' '+word_texts[k]
    de_word = [word.strip().strip(',') for word in de_word]
    en_word = [word.strip().strip(',') for word in en_word]
    vi_word = [word.strip().strip(',') for word in vi_word]
    type_word = [word.strip().strip(',') for word in type_word]
    field_word = [word.strip().strip(',') for word in field_word]
    out_message = 'Line has been analyzed successfully to ' + str(item+1) + ' item(s).'
    return de_word, type_word, field_word, en_word, vi_word, out_message

# Tests:
line = doc.paragraphs[2]
line = doc.paragraphs[16]
line = doc.paragraphs[19]
line = doc.paragraphs[25]

#for phrase in line.runs:
#   print(phrase.text)
#   print(line.style.name)
#   print(line.style.font.bold)
#   print(line.style.font.italic)
#   print(phrase.style.name)
#   print(phrase.style.font.bold)
#   print(phrase.style.font.italic)

paragraph_style_bold = line.style.font.bold
paragraph_style_italic = line.style.font.italic
word_texts = [part.text.strip() for part in line.runs if part.text.strip() != '']
word_style_names = [part.style.name for part in line.runs if part.text.strip() != '']
word_format_bolds, word_format_italics = read_format(word_style_names, paragraph_style_bold, paragraph_style_italic)
#word_format_bolds = [part.style.font.bold for part in line.runs if part.text.strip() != '']
#word_format_italics = [part.style.font.italic for part in line.runs if part.text.strip() != '']

for k in range(len(word_style_names)):
    print(word_texts[k])
    print(word_style_names[k])
    print(word_format_bolds[k])
    print(word_format_italics[k])

newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics, newlinerole_comment = re_parse(word_texts, word_format_bolds, word_format_italics)

de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics, newlinerole_comment)

# Main operation:

de_words = []
en_words = []
vi_words = []
type_words = []
field_words = []
out_messages = []

for line in doc.paragraphs:
    paragraph_style_bold = line.style.font.bold
    paragraph_style_italic = line.style.font.italic
    word_texts = [part.text.strip() for part in line.runs if part.text.strip() != '']
    word_style_names = [part.style.name for part in line.runs if part.text.strip() != '']
    word_format_bolds, word_format_italics = read_format(word_style_names, paragraph_style_bold, paragraph_style_italic)
    print(word_texts)
    newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics, newlinerole_comment = re_parse(word_texts, word_format_bolds, word_format_italics)
    print(newword_texts)    
    de_word, type_word, field_word, en_word, vi_word, out_message = analyze_line(newword_texts, newwordcase_capitals, newword_format_bolds, newword_format_italics, newlinerole_comment)
    print(de_word)
    print(en_word)
    print(vi_word)
    de_words.extend(de_word)
    en_words.extend(en_word)
    vi_words.extend(vi_word)
    type_words.extend(type_word)
    field_words.extend(field_word)
    out_messages.append(out_message)
    
# Note: line.runs[0].style.font.bold and line.runs[0].style.font.italic meaning:
    # line.runs[0].style.font.bold is "toggle" for line.style.font.bold, i.e.
    # if line.style.font.bold == True, and line.runs[0].style.font.bold == True
    # then the real style = False. Probably the XOR operator of the two booleans works. 
    
# Later: output directly the result from each line to the CSV file
